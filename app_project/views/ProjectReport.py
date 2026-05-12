import io
from datetime import datetime
from django.http import HttpResponse
from django.views import View
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app_project.models import Project
from app_project.mixins import ProjectAccessMixin

class ProjectReportExportView(ProjectAccessMixin, View):
    """导出项目进度报告：需有 view_project 权限"""
    permission_required = 'app_project.view_project'

    def get(self, request, pk):
        # 1. 获取项目对象并检查权限
        project = get_object_or_404(Project.objects.select_related(
            'manager',
            'repository',
            'repository__customer',
            'repository__oem',
            'material',
            'repository__salesperson'
        ).prefetch_related('nodes'), pk=pk)
        
        self.check_object_permission(project)

        # 2. 创建 Word 文档
        document = Document()
        
        # 设置中文字体兼容
        document.styles['Normal'].font.name = u'微软雅黑'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # --- 封面 ---
        self._add_cover(document, project)
        document.add_page_break()

        # --- 各章节内容 ---
        self._add_chapter_overview(document, project)
        self._add_chapter_business(document, project)
        self._add_chapter_material(document, project)
        self._add_chapter_progress(document, project)

        # 3. 输出文件流
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = f"项目进度报告_{project.name}_{datetime.now().strftime('%Y%m%d')}.docx"
        import urllib.parse
        filename = urllib.parse.quote(filename)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    def _add_heading(self, document, text, level=1):
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_paragraph(self, document, text, bold=False):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        if bold: run.font.bold = True
        return p

    def _add_cover(self, document, project):
        for _ in range(5): document.add_paragraph()
        title = document.add_heading(project.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            run.font.size = Pt(26)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)

        subtitle = document.add_paragraph("项目进度汇报报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.runs[0]
        run.font.size = Pt(18)
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        for _ in range(8): document.add_paragraph()
        info = document.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_text = f"项目负责人：{project.manager.username}\n生成日期：{datetime.now().strftime('%Y-%m-%d')}"
        run = info.add_run(info_text)
        run.font.size = Pt(12)
        run.font.name = u'微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    def _add_chapter_overview(self, document, project):
        self._add_heading(document, "1. 项目概况")
        table = document.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "项目名称"
        table.rows[0].cells[1].text = project.name
        table.rows[1].cells[0].text = "当前阶段"
        table.rows[1].cells[1].text = project.get_current_stage_display()
        table.rows[2].cells[0].text = "总体进度"
        table.rows[2].cells[1].text = f"{project.progress_percent}%"
        self._add_heading(document, "项目背景与描述", level=2)
        self._add_paragraph(document, project.description or "暂无描述")

    def _add_chapter_business(self, document, project):
        self._add_heading(document, "2. 商业与产品信息")
        repo = getattr(project, 'repository', None)
        if not repo:
            self._add_paragraph(document, "暂无档案信息")
            return
        table = document.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "直接客户"
        table.rows[0].cells[1].text = repo.customer.company_name if repo.customer else "-"
        table.rows[0].cells[2].text = "终端主机厂"
        table.rows[0].cells[3].text = repo.oem.name if repo.oem else "-"
        table.rows[1].cells[0].text = "产品名称"
        table.rows[1].cells[1].text = repo.product_name or "-"
        table.rows[1].cells[2].text = "产品代码"
        table.rows[1].cells[3].text = repo.product_code or "-"
        table.rows[2].cells[0].text = "目标成本"
        table.rows[2].cells[1].text = f"¥{repo.target_cost}" if repo.target_cost else "-"
        table.rows[2].cells[2].text = "竞品售价"
        table.rows[2].cells[3].text = f"¥{repo.competitor_price}" if repo.competitor_price else "-"
        table.rows[3].cells[0].text = "跟进业务员"
        # 兼容自定义 User 属性
        sp_name = repo.salesperson.get_full_name() or repo.salesperson.username if repo.salesperson else "-"
        table.rows[3].cells[1].text = sp_name
        table.rows[3].cells[2].text = "联系电话"
        table.rows[3].cells[3].text = repo.salesperson.phone if repo.salesperson else "-"

    def _add_chapter_material(self, document, project):
        self._add_heading(document, "3. 材料方案")
        if not project.material:
            self._add_paragraph(document, "暂未选定材料")
            return
        mat = project.material
        self._add_paragraph(document, f"选定材料：{mat.grade_name} ({mat.manufacturer})", bold=True)
        self._add_paragraph(document, f"材料类型：{mat.category.name} | 阻燃等级：{mat.flammability}")
        self._add_heading(document, "核心性能指标", level=2)
        grouped_props = mat.get_grouped_properties()
        if not grouped_props:
            self._add_paragraph(document, "暂无性能数据")
        else:
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "分类", "指标名称", "测试标准", "数值"
            for group in grouped_props:
                for item in group['items']:
                    row = table.add_row().cells
                    row[0].text = group['category_name']
                    row[1].text = item['name']
                    std = item['standard']
                    if item['condition']: std += f" ({item['condition']})"
                    row[2].text = std
                    val = str(item['value'])
                    if item['unit']: val += f" {item['unit']}"
                    row[3].text = val

    def _add_chapter_progress(self, document, project):
        self._add_heading(document, "4. 项目进度详情")
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "阶段", "状态", "更新时间", "备注/进展"
        for node in project.cached_nodes:
            row = table.add_row().cells
            name = node.get_stage_display()
            if node.round > 1: name += f" (第{node.round}轮)"
            row[0].text = name
            row[1].text = node.get_status_display()
            row[2].text = node.updated_at.strftime('%Y-%m-%d') if node.status != 'PENDING' else "-"
            row[3].text = node.remark or "-"
