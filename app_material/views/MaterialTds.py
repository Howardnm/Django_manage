import io
from datetime import datetime
from urllib.parse import quote

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.views import View

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from init import TDS_TEMPLATE_PATH
from app_material.models.material import MaterialLibrary
from app_material.mixins import MaterialAccessMixin

# 分类英文名映射（示例 TDS 排版用；未知分类只显示中文）
CATEGORY_EN = {
    '物理性能': 'Physical Properties',
    '机械性能': 'Mechanical Properties',
    '热学性能': 'Thermal Properties',
    '阻燃/电气': 'Flammability / Electrical',
    '老化/耐候': 'Aging / Weathering',
    '其他性能': 'Others Properties',
}


# 加工条件表（Table 1）固定行序 R1-R8 对应的模型字段组（典型值, 范围）
PROCESSING_ROW_FIELDS = [
    ('melt_temp_value', 'melt_temp_range'),
    ('barrel_rear_value', 'barrel_rear_range'),
    ('barrel_center_value', 'barrel_center_range'),
    ('barrel_front_value', 'barrel_front_range'),
    ('mold_temp_value', 'mold_temp_range'),
    ('temp_limit_value', 'temp_limit_range'),
    ('injection_speed_value', 'injection_speed_range'),
    ('pre_dry_value', 'pre_dry_range'),
]


def _fill_processing_table(table, material):
    """填充加工条件表（Table 1）：按模板固定行序 R1-R8 写入「典型值」「范围」两列。

    无数据写空字符串——清空模板示例占位值（220℃ 等），避免误导。
    模板 R6-R8 的典型值/范围两列是合并单元格（cells[2] is cells[3]），只写一次。
    若模板行数多于 8（用户人工增行），多出的行保留模板原样。
    """
    pc = getattr(material, 'processing_condition', None)
    for row, (val_field, range_field) in zip(table.rows[1:], PROCESSING_ROW_FIELDS):
        val = getattr(pc, val_field, '') if pc else ''
        rng = getattr(pc, range_field, '') if pc else ''
        cells = row.cells
        if cells[2] is cells[3]:
            # 合并单元格：只写一次（优先典型值，典型值空则写范围）
            _write_cell(cells[2], val or rng)
        else:
            _write_cell(cells[2], val)
            _write_cell(cells[3], rng)


def _replace_para_text(p, text):
    """替换段落文本：保留含图形（w:drawing，如蓝色色带）的 run，仅收拢文本 run 改写。

    注意：不能对含图形的 run 直接执行 run.text 赋值，python-docx 会清空该 run 内所有子元素（含图形）。
    """
    text_runs = [r for r in p.runs if r._element.find(qn('w:t')) is not None]
    if not text_runs:
        # 没有文本 run（如图形 run），直接追加一个文本 run
        p.add_run(text)
        return
    text_runs[0].text = text
    for r in text_runs[1:]:
        r._element.getparent().remove(r._element)


def _fmt_num(value):
    """数值去尾零显示，如 17 / 1.072 / 0.7"""
    if value is None:
        return ''
    return '%g' % float(value)


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def _write_cell(cell, text, *, bold=False, white=False, left=False, size=9):
    """向单元格写入文本并应用与模板一致的字体（Times New Roman + 微软雅黑）。"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if not left:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _build_property_table(table, grouped_properties, value_mode='range'):
    """重建物性表：保留表头，删除正文行，按分类分组写入。

    典型值规则（NUMBER）：
      range 模式（默认）：min_value 与 max_value 都存在 → 优先写范围值；min == max 时只显示单值；
                           仅 min 存在 → 写 ≥ min；仅 max 存在 → 写 ≤ max；
                           两者皆无 → 退化写 value；
      single 模式：只写默认数值 value（不合成范围）；
    TEXT/SELECT 用 value_text。
    """
    # 1. 删除除表头外的所有正文行
    body_trs = [row._tr for row in table.rows[1:]]
    for tr in body_trs:
        tr.getparent().remove(tr)

    # 2. 重建
    for group in grouped_properties:
        cat_name = group['category_name']
        # 分类行：5 列合并，蓝色底纹 + 白色加粗
        cat_row = table.add_row()
        cat_cell = cat_row.cells[0].merge(cat_row.cells[4])
        _set_cell_shading(cat_cell, '4166F5')
        cat_en = CATEGORY_EN.get(cat_name, '')
        cat_text = f"{cat_name} {cat_en}".strip()
        _write_cell(cat_cell, cat_text, bold=True, white=True)

        # 指标行
        for item in group['items']:
            row = table.add_row()
            name_en = item.get('name_en', '')
            prop_name = f"{item['name']} {name_en}".strip() if name_en else item['name']
            _write_cell(row.cells[0], prop_name, left=True)

            _write_cell(row.cells[1], item['standard'] or '')
            _write_cell(row.cells[2], item['condition'] or '')
            _write_cell(row.cells[3], item['unit'] or '')

            # 典型值
            if item['data_type'] == 'NUMBER':
                if value_mode == 'single':
                    val_text = _fmt_num(item['value'])
                else:
                    min_v, max_v = item['min_value'], item['max_value']
                    if min_v is not None and max_v is not None:
                        val_text = _fmt_num(min_v) if min_v == max_v else f"{_fmt_num(min_v)} ~ {_fmt_num(max_v)}"
                    elif min_v is not None:
                        val_text = f"≥ {_fmt_num(min_v)}"
                    elif max_v is not None:
                        val_text = f"≤ {_fmt_num(max_v)}"
                    else:
                        val_text = _fmt_num(item['value'])
            else:
                # get_grouped_properties 对 TEXT/SELECT 已把 value_text 存入 value 键
                val_text = item['value'] or ''
            _write_cell(row.cells[4], val_text)

    # 无任何数据时给出占位提示
    if not grouped_properties:
        row = table.add_row()
        cell = row.cells[0].merge(row.cells[4])
        _write_cell(cell, '暂无性能数据', bold=True)


class MaterialTdsExportView(MaterialAccessMixin, View):
    """导出材料物性表 (TDS) docx：需具备查看权限。"""
    permission_required = 'app_material.view_materiallibrary'

    def get(self, request, pk):
        material = get_object_or_404(
            MaterialLibrary.objects.select_related('category', 'processing_condition').prefetch_related(
                'characteristics', 'scenarios', 'properties__test_config__category'
            ),
            pk=pk,
        )
        self.check_object_permission(material)

        document = Document(str(TDS_TEMPLATE_PATH))

        # 1. 标题：所有「牌号 …」开头的段落（含第二页重复标题）替换为实际牌号（不含材料类型）
        title = f"牌号 {material.grade_name}"
        for p in document.paragraphs:
            if p.text.strip().startswith('牌号'):
                _replace_para_text(p, title)

        # 2. 主要特征 / 主要应用
        features = '、'.join(c.name for c in material.characteristics.all()) or '—'
        applications = '、'.join(s.name for s in material.scenarios.all()) or '—'
        for p in document.paragraphs:
            t = p.text.strip()
            if t.startswith('主要特征'):
                _replace_para_text(p, f"主要特征：{features}")
            elif t.startswith('主要应用'):
                _replace_para_text(p, f"主要应用：{applications}")

        # 3. 物性表（Table 0）动态重建；加工条件表（Table 1）按固定 8 行填充
        #    典型值策略：range=优先范围值（默认），single=只填默认数值
        value_mode = self.request.GET.get('value_mode', 'range')
        _build_property_table(document.tables[0], material.get_grouped_properties(), value_mode)
        _fill_processing_table(document.tables[1], material)

        # 4. 输出文件流
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = f"{material.grade_name}_TDS_{datetime.now().strftime('%Y%m%d')}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
        return response